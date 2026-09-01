import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_manual():
    doc = Document()
    
    # Page setup - Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Styles
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Segoe UI'
    style_normal.font.size = Pt(10.5)
    style_normal.font.color.rgb = RGBColor(40, 44, 52)
    
    # Title Page / Header Block
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("LEADS NEXT GEN CENTRE\nOPERATIONS & PRIVILEGES MANUAL")
    title_run.font.name = 'Segoe UI'
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(24, 75, 120)
    
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("Comprehensive System Instruction Guide, Role Privilege Matrix & Security Architecture")
    sub_run.font.size = Pt(12)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(100, 110, 125)
    
    doc.add_paragraph()
    
    # Metadata Box
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Platform:", "LEADS ERP All-in-One Operations Portal"),
        ("Institution:", "LEADS Next Gen Centre, M.S. Ramaiah University of Applied Sciences (MSRUAS)"),
        ("System Version:", "v2.0 Production (Next.js 16 + AES-256-GCM Secure Data Layer)"),
        ("Document Classification:", "Internal Operational & Administrative Reference"),
    ]
    for i, (k, v) in enumerate(meta_data):
        row = meta_table.rows[i]
        c1, c2 = row.cells[0], row.cells[1]
        c1.width = Inches(2.2)
        c2.width = Inches(4.3)
        set_cell_background(c1, "F0F4F8")
        set_cell_background(c2, "FAFAFA")
        set_cell_margins(c1, 80, 80, 120, 120)
        set_cell_margins(c2, 80, 80, 120, 120)
        
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(k)
        r1.bold = True
        r1.font.size = Pt(10)
        r1.font.color.rgb = RGBColor(24, 75, 120)
        
        p2 = c2.paragraphs[0]
        r2 = p2.add_run(v)
        r2.font.size = Pt(10)
        
    doc.add_paragraph()
    doc.add_heading("1. Executive System Overview", level=1)
    doc.add_paragraph(
        "The LEADS Next Gen ERP Dashboard is a centralized institutional management platform built specifically "
        "for the LEADS Next Gen Centre at MSRUAS. The platform manages events, real-time task allocations, automated "
        "reimbursement verification pipelines, budget allocation tracking, design asset clearance workflows, dynamic public feedback "
        "forms, and hierarchical organizational member directories."
    )
    
    doc.add_heading("2. One-Time Initial Setup Wizard", level=1)
    doc.add_paragraph(
        "Upon deploying a fresh instance of the LEADS ERP (locally or on a production VPS), the application automatically enters "
        "the Initial Setup Wizard upon first launch. Once completed, this setup is permanently locked and cannot be re-executed."
    )
    
    doc.add_heading("Step 1: Super User Account Provisioning", level=2)
    doc.add_paragraph(
        "• The operator sets up the primary root Super User (Name, Email ID, and Master Password min 8 chars).\n"
        "• Scrypt cryptographic hashing (with unique salts and timing-safe comparisons) is used to protect credentials.\n"
        "• A fresh install starts with zero pre-existing accounts; all additional faculty, staff, and student accounts are created via the Members Directory."
    )
    
    doc.add_heading("Step 2: Database Server-Side Encryption Key Setup", level=2)
    doc.add_paragraph(
        "• The system prompts the operator to configure the DATA_ENCRYPTION_KEY used to encrypt all local database files on disk.\n"
        "• Operators can either auto-generate a cryptographically secure 256-bit key (recommended) or specify a custom secret key.\n"
        "• The key is saved permanently to .env on the server. The setup wizard displays a prominent alert reminding the administrator to store a secure offline backup of the key.\n"
        "• All local database collections (data/members.json, data/events.json, etc.) are encrypted at rest using AES-256-GCM authenticated encryption."
    )
    
    doc.add_heading("3. Granular Role & Privilege Matrix (Tiers 1 to 7)", level=1)
    
    tier_table = doc.add_table(rows=1, cols=4)
    tier_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tier_table.rows[0].cells
    hdr[0].width = Inches(1.0)
    hdr[1].width = Inches(1.5)
    hdr[2].width = Inches(1.5)
    hdr[3].width = Inches(2.5)
    
    headers = ["Tier", "Role Name", "Division", "Key System Permissions & Scope"]
    for i, h in enumerate(headers):
        set_cell_background(hdr[i], "184B78")
        set_cell_margins(hdr[i], 100, 100, 100, 100)
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(10)
        
    tiers_data = [
        ("Tier 1", "Super User", "Core Committee", "Unrestricted system control, global audit logs, emergency lockdown mode, policy management, backup/restore, dynamic Quick Switch impersonation."),
        ("Tier 2", "Centre Head", "Faculty", "Full operational authority across both campuses, final budget approval, Level-2 reimbursement clearance, guest directory management, email broadcasts."),
        ("Tier 2.5", "GG Campus Head", "Faculty", "Campus-specific oversight for Gnanagangothri (GG) campus operations, events, and task coordination."),
        ("Tier 3", "Faculty / Event Head / Advisors", "Faculty", "Event proposal approvals, Level-1 reimbursement audits, task assignment to student leads, performance evaluation reviewer."),
        ("Tier 4", "Advisory Board", "Faculty", "High-level read access to institutional performance metrics, reports, ratings, and event summaries."),
        ("Tier 5", "Core Committee (Officers & Leads)", "Core Committee", "Operational execution, task allocation, event planning, form building & public QR code publishing, financial claim submission."),
        ("Tier 6", "Training Associate / Members", "Training Associate", "Personal workspace access, assigned task execution & status updates, reimbursement claim submission, feedback submission."),
        ("Tier 7", "Alumni / Guests", "Alumni / Guest", "View-only access to relevant past archives, certificates, or assigned guest directories."),
    ]
    
    for row_data in tiers_data:
        row = tier_table.add_row().cells
        row[0].width = Inches(1.0)
        row[1].width = Inches(1.5)
        row[2].width = Inches(1.5)
        row[3].width = Inches(2.5)
        for i, val in enumerate(row_data):
            set_cell_background(row[i], "FFFFFF" if int(row_data[0].split()[1].replace('.5','')) % 2 == 0 else "F9FBFD")
            set_cell_margins(row[i], 80, 80, 100, 100)
            p = row[i].paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(9.5)
            if i == 0:
                r.bold = True
                
    doc.add_paragraph()
    doc.add_heading("4. Comprehensive Module Operating Instructions", level=1)
    
    modules = [
        ("4.1 Members Directory & User Management",
         "The Members Directory allows administrators to add, update, terminate, and reactivate members. "
         "When adding a member, an activation link with a secure cryptographic token is generated. "
         "Super Users can also trigger password resets, require mandatory password updates on next login, "
         "or assign members to specific divisions and departments."),
        
        ("4.2 Dynamic Quick Switch (Super User Feature)",
         "The Quick Switch feature (accessible via the user settings header) allows the Super User to instantly view "
         "the dashboard from the perspective of any registered account without requiring their password. "
         "The switcher dynamically queries the live member roster, ensuring only real, active accounts appear. "
         "A prominent top bar allows one-click return to the Super User session at any time."),
        
        ("4.3 Events & Automated Task Delegation",
         "Events progress through draft, proposal, faculty approval, active execution, and archival stages. "
         "Sub-committees (Stage, Logistics, Hospitality, Social Media) can be attached to events with dedicated "
         "lead allocations. Tasks created within events automatically update progress percentages and trigger audit logs."),
        
        ("4.4 Design Portal & Multi-Gate Proofreading Engine",
         "The Design Portal manages promotional creatives and posters. Uploaded assets undergo dual verification: "
         "Gate 1 (Style & Resolution check) and Gate 2 (AI OCR Spellcheck & designated Faculty Proofreader sign-off). "
         "Once approved, associated design deliverables automatically mark linked event tasks as complete."),
        
        ("4.5 Multi-Stage Financial Reimbursement Pipeline",
         "Claims submitted by members include digitized receipts and event tags. The pipeline enforces dual-approval: "
         "Level 1 Faculty Verification followed by Level 2 Centre Head Clearance. Disbursed funds automatically adjust "
         "event budgets and sponsor balances, calculating net university expenditure in real time."),
        
        ("4.6 Dynamic Public Forms & Feedback System",
         "Administrators can create public sign-up and feedback forms with custom fields (text, dropdowns, ratings). "
         "Each form generates a dynamic QR code and shareable URL. Responses are logged into the encrypted database "
         "with export capabilities to CSV, PDF, and Word DOCX summary reports."),
        
        ("4.7 Email Management & Broadcast Communication",
         "Outbound announcements and notifications route through a local Postfix relay authenticated with MSRUAS Workspace. "
         "Custom scopes (All Members, Core Committee, Faculty Only) allow targeted communication with delivery tracking.")
    ]
    
    for title, desc in modules:
        doc.add_heading(title, level=2)
        doc.add_paragraph(desc)
        
    doc.add_heading("5. Security, Backup & Maintenance Guidelines", level=1)
    doc.add_paragraph(
        "• Data Encryption: Never commit the DATA_ENCRYPTION_KEY to git repositories. It must remain strictly in the .env file.\n"
        "• Backups: Use the built-in Backup & Restore module (Tier 1 Super User) to generate encrypted snapshots. Backups can be decrypted offline using scripts/decrypt-backup.js with the master key.\n"
        "• Deployment on VPS: Follow the standard deployment sequence (git pull -> npm install -> npm run build -> pm2 restart leads-dashboard).\n"
        "• Disaster Recovery: If migrating servers, copy both the data/ directory and the DATA_ENCRYPTION_KEY string from .env."
    )
    
    doc.add_paragraph()
    p_end = doc.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_end = p_end.add_run("— End of Operations Manual —\n© 2026 LEADS Next Gen Centre, MSRUAS. All rights reserved.")
    r_end.font.size = Pt(9)
    r_end.font.italic = True
    r_end.font.color.rgb = RGBColor(120, 120, 120)
    
    os.makedirs("docs", exist_ok=True)
    target_path = os.path.join(os.getcwd(), "docs", "LEADS_ERP_Instruction_and_Privileges_Manual.docx")
    doc.save(target_path)
    print(f"Successfully generated manual at: {target_path}")

if __name__ == "__main__":
    create_manual()
